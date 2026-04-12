import json
from pathlib import Path
from typing import Awaitable, Callable, Optional
import PyPDF2
from app.core.ai_client import chat

async def extract_text_from_pdf(filepath: Path) -> str:
    text = ""
    with open(filepath, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text.strip()


async def extract_text_from_docx(filepath: Path) -> str:
    """Extract text from a Microsoft Word .docx file (paragraphs + tables)."""
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from e

    doc = Document(str(filepath))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # Pull text out of tables as well — résumés often live in tables.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()


async def extract_text_from_resume_file(filepath: Path, filename: str = "") -> str:
    """
    Extract plain text from a résumé file. Supports PDF (.pdf) and Word (.docx).

    Args:
        filepath: path to the temp file on disk
        filename: original upload filename (used to detect extension)
    """
    name = (filename or filepath.name or "").lower()
    if name.endswith(".pdf"):
        return await extract_text_from_pdf(filepath)
    if name.endswith(".docx"):
        return await extract_text_from_docx(filepath)
    if name.endswith(".doc"):
        raise ValueError(
            "Legacy .doc files are not supported — please save as .docx or .pdf and try again."
        )
    raise ValueError("Unsupported file type. Upload a PDF or Word (.docx) résumé.")


async def parse_resume_file(
    filepath: Path,
    progress_callback: Optional[Callable[[str], Awaitable[None]]] = None,
) -> dict:
    """Parse a résumé file (PDF or .docx) into the structured profile dict."""
    if progress_callback:
        await progress_callback("extracting_text")
    raw_text = await extract_text_from_resume_file(filepath, filepath.name)
    if not raw_text:
        return _empty_parse()
    return await _parse_resume_text(raw_text, progress_callback)


async def parse_resume_pdf(
    filepath: Path,
    progress_callback: Optional[Callable[[str], Awaitable[None]]] = None
) -> dict:
    if progress_callback:
        await progress_callback("extracting_text")
    raw_text = await extract_text_from_pdf(filepath)
    if not raw_text:
        return _empty_parse()
    return await _parse_resume_text(raw_text, progress_callback)


async def _parse_resume_text(
    raw_text: str,
    progress_callback: Optional[Callable[[str], Awaitable[None]]] = None,
) -> dict:

    prompt = f"""You are a resume parser. Extract structured information from this resume text.
Return ONLY valid JSON with these exact keys:
{{
  "skills": ["skill1", "skill2"],
  "titles": ["Most Recent Title", "Previous Title"],
  "years_experience": 10,
  "summary": "2-3 sentence professional summary",
  "role_family": "SRE | DevOps | Software Engineering | Data Engineering | Product Management | Design | Finance | Marketing | Sales | General"
}}

Rules:
- skills: up to 20 most relevant technical and soft skills
- titles: job titles from the resume, most recent first
- years_experience: total years as a number
- role_family: pick the single best match from the options above
- summary: write in third person, highlight seniority and key strengths

Resume text:
{raw_text[:4000]}
"""

    if progress_callback:
        await progress_callback("sending_to_llm")
    try:
        result = await chat([{"role": "user", "content": prompt}], json_mode=True)
    except Exception as e:
        # Surface LLM failures (Ollama not running, OpenAI key invalid, etc.)
        # so the upload pipeline can mark the profile as "failed" with a clear
        # message instead of silently producing an empty "done" profile.
        raise RuntimeError(
            f"LLM call failed while parsing résumé. Check that your AI provider is running "
            f"and configured in Settings. ({type(e).__name__}: {e})"
        ) from e

    if progress_callback:
        await progress_callback("classifying")
    try:
        parsed = json.loads(result)
    except json.JSONDecodeError as e:
        raise RuntimeError(
            f"LLM returned invalid JSON. The model may be too small for structured output — "
            f"try a larger model in Settings. ({e})"
        ) from e

    parsed["raw_text"] = raw_text
    parsed["skills"] = json.dumps(parsed.get("skills", []))
    parsed["titles"] = json.dumps(parsed.get("titles", []))
    return parsed

def _empty_parse(raw_text: str = "") -> dict:
    return {
        "raw_text": raw_text,
        "skills": "[]",
        "titles": "[]",
        "years_experience": 0,
        "summary": "",
        "role_family": "General",
    }
